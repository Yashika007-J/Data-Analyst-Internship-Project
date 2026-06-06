from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title Page
title = doc.add_heading('Care Transition Efficiency &', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

title2 = doc.add_heading('Placement Outcome Analytics', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Submitted by: Yashika \n').bold = True
info.add_run('Internship Program: Unified Mentor — Data Analytics\n').bold = True
info.add_run('Duration: 2025–2026\n').bold = True
info.add_run('Dataset: U.S. HHS UAC Program Data\n').bold = True

doc.add_page_break()

# 1. Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report presents the findings of a data analytics project conducted '
    'as part of the Unified Mentor Data Analytics Internship. The project '
    'focused on analyzing the efficiency of the U.S. Unaccompanied Alien '
    'Children (UAC) care pipeline using 720 days of operational data spanning '
    'January 2023 to December 2025. Custom KPIs were derived, exploratory '
    'analysis was performed, and an interactive Streamlit dashboard was built '
    'to present insights to stakeholders.'
)

# 2. Problem Statement
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Existing government reporting focuses only on how many children are '
    'in custody at a given time. This project identified a critical gap — '
    'the absence of process efficiency metrics. Without these, bottlenecks '
    'and delays in child placement remain hidden and unaddressed.'
)
doc.add_paragraph('Key questions addressed:', style='Normal')
doc.add_paragraph('How quickly are children moved from CBP to HHS care?', style='List Bullet')
doc.add_paragraph('Are daily discharges keeping pace with daily arrivals?', style='List Bullet')
doc.add_paragraph('Where and when do care backlogs accumulate?', style='List Bullet')
doc.add_paragraph('Are placement outcomes improving over time?', style='List Bullet')

# 3. Dataset
doc.add_heading('3. Dataset Description', level=1)
doc.add_paragraph('Source: U.S. Department of Health and Human Services (HHS)', style='List Bullet')
doc.add_paragraph('Time Period: January 2023 – December 2025', style='List Bullet')
doc.add_paragraph('Total Records: 720 daily observations', style='List Bullet')
doc.add_paragraph('Format: CSV — 6 columns', style='List Bullet')

doc.add_paragraph('')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Column'
hdr[1].text = 'Description'
rows = [
    ('Date', 'Reporting date'),
    ('CBP Apprehended', 'Daily children entering CBP custody'),
    ('CBP Custody', 'Total children in CBP care'),
    ('CBP Transferred', 'Children moved to HHS'),
    ('HHS Care', 'Total children in HHS care'),
    ('HHS Discharged', 'Children placed with sponsors'),
]
for col, desc in rows:
    row = table.add_row().cells
    row[0].text = col
    row[1].text = desc

# 4. Methodology
doc.add_heading('4. Methodology', level=1)
doc.add_paragraph('Data Cleaning — removed nulls, fixed date formats, standardized column names', style='List Number')
doc.add_paragraph('EDA — plotted all 5 columns as line charts to understand trends', style='List Number')
doc.add_paragraph('KPI Derivation — calculated 3 custom efficiency metrics', style='List Number')
doc.add_paragraph('Dashboard — built interactive Streamlit app with date filters', style='List Number')
doc.add_paragraph('Report — documented all findings in Word format', style='List Number')

# 5. KPIs
doc.add_heading('5. KPIs Calculated', level=1)
doc.add_paragraph('Transfer Efficiency Ratio = CBP Transferred ÷ CBP Apprehended', style='List Bullet')
doc.add_paragraph('Discharge Effectiveness = HHS Discharged ÷ HHS Care Load', style='List Bullet')
doc.add_paragraph('Pipeline Throughput = (Transferred + Discharged) ÷ (Apprehended + Transferred)', style='List Bullet')

# 6. Key Findings
doc.add_heading('6. Key Findings', level=1)
doc.add_paragraph('HHS Care peaked at 11,516 children in December 2023', style='List Bullet')
doc.add_paragraph('Sharp decline in arrivals after January 2025 — possible policy change', style='List Bullet')
doc.add_paragraph('Average Transfer Efficiency = 1.50 — system was clearing backlog', style='List Bullet')
doc.add_paragraph('Average Discharge Rate = 0.02 — sponsor placement remained slow', style='List Bullet')
doc.add_paragraph('Weekday transitions faster than weekends', style='List Bullet')
doc.add_paragraph('Longest backlog period: August 2023 to January 2024', style='List Bullet')

# 7. Tools Used
doc.add_heading('7. Tools & Technologies Used', level=1)
doc.add_paragraph('Python 3.14', style='List Bullet')
doc.add_paragraph('Pandas — data cleaning and manipulation', style='List Bullet')
doc.add_paragraph('Matplotlib & Seaborn — data visualization', style='List Bullet')
doc.add_paragraph('Streamlit — interactive dashboard', style='List Bullet')
doc.add_paragraph('VS Code — development environment', style='List Bullet')
doc.add_paragraph('Jupyter Notebook — EDA and analysis', style='List Bullet')

# 8. Conclusion
doc.add_heading('8. Conclusion', level=1)
doc.add_paragraph(
    'This project successfully transformed raw government data into meaningful '
    'efficiency insights. The KPIs and dashboard developed here provide a '
    'practical framework for monitoring child welfare pipeline performance. '
    'The findings highlight the need for structured transition analytics '
    'in government care systems and demonstrate how data science can '
    'support better humanitarian outcomes.'
)

# 9. References
doc.add_heading('9. References', level=1)
doc.add_paragraph('U.S. Department of Health and Human Services — ORR UAC Dataset', style='List Bullet')
doc.add_paragraph('Unified Mentor Data Analytics Internship, 2026', style='List Bullet')
doc.add_paragraph('Python Pandas Documentation — pandas.pydata.org', style='List Bullet')
doc.add_paragraph('Streamlit Documentation — docs.streamlit.io', style='List Bullet')

doc.save('UAC_Project_Report.docx')
print("Done! UAC_Project_Report.docx saved!")