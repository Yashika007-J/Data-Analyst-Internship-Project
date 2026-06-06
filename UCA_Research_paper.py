from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Care Transition Efficiency & Placement Outcome Analytics', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Research Paper — Unified Mentor Data Analytics Internship')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(13)

author = doc.add_paragraph('Author: Yashika Jandaniya | Data Analyst Intern | Year: 2026')
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'This research paper presents a comprehensive data analytics study of the '
    'Unaccompanied Alien Children (UAC) care pipeline in the United States. '
    'Using daily operational data sourced from the U.S. Department of Health '
    'and Human Services (HHS), this study examines process efficiency metrics '
    'across three stages: CBP apprehension, HHS care placement, and sponsor '
    'reunification. Key Performance Indicators (KPIs) including Transfer '
    'Efficiency Ratio, Discharge Effectiveness Index, and Pipeline Throughput '
    'Rate were derived and analyzed over a period spanning January 2023 to '
    'December 2025. Findings reveal significant systemic shifts in care load, '
    'transfer speed, and discharge patterns, with actionable insights presented '
    'through an interactive Streamlit dashboard.'
)

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'The UAC Program, administered jointly by U.S. Customs and Border Protection (CBP) '
    'and the Department of Health and Human Services (HHS), represents one of the most '
    'operationally complex child welfare systems in the United States. Each year, tens of '
    'thousands of unaccompanied minors enter the care pipeline, requiring rapid medical '
    'screening, shelter placement, case management, and eventual reunification with a '
    'vetted sponsor.'
)
doc.add_paragraph(
    'From a policy and humanitarian perspective, the speed, continuity, and reliability '
    'of this pipeline are as critical as raw capacity. Yet existing public reporting '
    'focuses primarily on aggregate custody counts, leaving process efficiency largely '
    'unmeasured. This study addresses that gap by applying data analytics techniques '
    'to derive actionable efficiency metrics from publicly available HHS datasets.'
)

# 2. Problem Statement
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'While aggregate counts of children in custody are monitored and reported, '
    'process efficiency metrics are largely absent from public discourse. '
    'Key unanswered questions include:'
)
bullets = [
    'How efficiently are children transferred from CBP to HHS custody?',
    'Are daily discharges keeping pace with daily inflows?',
    'When and where do care backlogs accumulate in the pipeline?',
    'Are placement outcomes improving or deteriorating over time?',
    'What temporal patterns exist in transition speed and success rates?'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Without structured transition analytics, systemic bottlenecks remain hidden, '
    'preventing timely policy intervention and resource allocation.'
)

# 3. Objectives
doc.add_heading('3. Project Objectives', level=1)
doc.add_heading('3.1 Primary Objectives', level=2)
primary = [
    'Measure the efficiency of CBP to HHS care transitions on a daily basis.',
    'Evaluate discharge and sponsor placement outcomes over time.',
    'Identify delays, bottlenecks, and periods of system stress.',
]
for p in primary:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('3.2 Secondary Objectives', level=2)
secondary = [
    'Support evidence-based recommendations for faster reunification.',
    'Improve case management workflows through data-driven insights.',
    'Inform policy-level process reforms using historical trend analysis.',
]
for s in secondary:
    doc.add_paragraph(s, style='List Bullet')

# 4. Dataset
doc.add_heading('4. Dataset Description', level=1)
doc.add_paragraph(
    'The dataset used in this study was sourced from the U.S. Department of '
    'Health and Human Services Office of Refugee Resettlement (ORR) and made '
    'available through the Unified Mentor internship platform.'
)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Column Name'
hdr[1].text = 'Description'

rows = [
    ('Date', 'Reporting date of the record'),
    ('Children Apprehended (CBP)', 'Daily intake volume into CBP custody'),
    ('Children in CBP Custody', 'Active CBP care load on that date'),
    ('Children Transferred out of CBP', 'Flow into HHS system from CBP'),
    ('Children in HHS Care', 'Active HHS care load on that date'),
    ('Children Discharged from HHS', 'Successful sponsor placements'),
]
for col, desc in rows:
    row = table.add_row().cells
    row[0].text = col
    row[1].text = desc

doc.add_paragraph('')
doc.add_paragraph('Total Records: 720 daily observations | Time Period: January 2023 – December 2025')

# 5. Methodology
doc.add_heading('5. Analytical Methodology', level=1)

doc.add_heading('5.1 Data Preprocessing', level=2)
doc.add_paragraph(
    'Raw data was loaded using the Python Pandas library. Column names were '
    'standardized, comma-formatted numeric values were cleaned, and date fields '
    'were converted to datetime format. Missing values were handled using forward '
    'fill and zero imputation strategies.'
)

doc.add_heading('5.2 Exploratory Data Analysis (EDA)', level=2)
doc.add_paragraph(
    'Five primary line charts were generated to understand individual column '
    'trends over time. A bar chart was used for monthly aggregation and a '
    'correlation heatmap was generated using Seaborn to identify relationships '
    'between variables.'
)

doc.add_heading('5.3 KPI Derivation', level=2)
kpis = [
    ('Transfer Efficiency Ratio', 'CBP Transferred ÷ CBP Apprehended', 'Measures speed of CBP to HHS transition'),
    ('Discharge Effectiveness Index', 'HHS Discharged ÷ HHS Care Load', 'Measures placement success rate'),
    ('Pipeline Throughput Rate', '(Transferred + Discharged) ÷ (Apprehended + Transferred)', 'Overall system movement'),
]
for name, formula, desc in kpis:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(f'{formula} — {desc}')

# 6. Key Findings
doc.add_heading('6. Key Findings', level=1)
findings = [
    'HHS Care load peaked in December 2023 at 11,516 children — the highest recorded value.',
    'Apprehensions dropped significantly after January 2025, indicating major policy shifts.',
    'Average Transfer Efficiency Ratio = 1.50 — system transferred more than it received daily.',
    'Average Discharge Effectiveness = 0.02 — discharge rates remained consistently low.',
    'Weekday transitions were significantly faster than weekend periods.',
    'Prolonged backlog periods were identified between August 2023 and January 2024.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

# 7. Conclusion
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    'This project successfully reframes the UAC dataset from a simple capacity '
    'monitoring lens to a comprehensive process efficiency and outcome evaluation '
    'framework. By deriving and analyzing KPIs such as Transfer Efficiency Ratio '
    'and Discharge Effectiveness Index, it provides actionable insights for '
    'improving reunification timelines, reducing systemic delays, and strengthening '
    'child welfare outcomes across the UAC care pipeline.'
)
doc.add_paragraph(
    'The interactive Streamlit dashboard developed as part of this project enables '
    'policymakers and case managers to explore efficiency trends dynamically, '
    'supporting evidence-based decision making in real time.'
)

# 8. References
doc.add_heading('8. References', level=1)
refs = [
    'Unified Mentor Data Analytics Internship Program, 2026',
    'Python Software Foundation — Pandas, Matplotlib, Seaborn Documentation',
    'Streamlit Inc. — Streamlit Framework Documentation',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.save('UAC_Research_Paper.docx')
print("Professional Research Paper saved!")